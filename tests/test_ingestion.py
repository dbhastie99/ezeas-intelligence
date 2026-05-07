from io import BytesIO

from docx import Document

from app.models.knowledge import KnowledgeChunk, KnowledgeDocument
from app.services.document_extraction_service import extract_text_from_docx
from app.services.ingestion_service import ingest_file_bytes


def test_txt_ingestion_creates_document_and_chunks(client, db_session):
    content = b"Minerva is advisory. It reads and audits knowledge evidence. " * 20

    response = client.post(
        "/api/v1/ingest/file",
        files={"file": ("doctrine.txt", content, "text/plain")},
        data={"source_type": "PLATFORM_DOCTRINE", "capability_status": "DOCTRINE"},
    )

    assert response.status_code == 200
    body = response.json()
    assert body["duplicate"] is False
    assert body["chunk_count"] > 0
    assert body["source_type"] == "PLATFORM_DOCTRINE"

    documents = db_session.query(KnowledgeDocument).all()
    chunks = db_session.query(KnowledgeChunk).all()
    assert len(documents) == 1
    assert len(chunks) == body["chunk_count"]
    assert documents[0].FileSha256 == body["file_sha256"]


def test_duplicate_ingestion_returns_duplicate_and_does_not_create_extra_chunks(client, db_session):
    content = b"Duplicate prevention by SHA256 keeps Minerva ingestion idempotent."

    first = client.post(
        "/api/v1/ingest/file",
        files={"file": ("duplicate.txt", content, "text/plain")},
        data={"source_type": "OTHER"},
    )
    second = client.post(
        "/api/v1/ingest/file",
        files={"file": ("duplicate-renamed.txt", content, "text/plain")},
        data={"source_type": "OTHER"},
    )

    assert first.status_code == 200
    assert second.status_code == 200
    assert second.json()["duplicate"] is True
    assert second.json()["document_id"] == first.json()["document_id"]
    assert db_session.query(KnowledgeDocument).count() == 1
    assert db_session.query(KnowledgeChunk).count() == first.json()["chunk_count"]


def test_duplicate_integrity_race_returns_duplicate_without_extra_chunks(db_session, monkeypatch):
    content = b"Duplicate race handling should survive a missed pre-insert SHA check."
    original_document, duplicate = ingest_file_bytes(
        db=db_session,
        content=content,
        original_file_name="race.txt",
        source_type="OTHER",
    )
    assert duplicate is False
    original_scalar = db_session.scalar
    calls = {"count": 0}

    def miss_first_duplicate_check(statement, *args, **kwargs):
        if calls["count"] == 0:
            calls["count"] += 1
            return None
        return original_scalar(statement, *args, **kwargs)

    monkeypatch.setattr(db_session, "scalar", miss_first_duplicate_check)

    duplicate_document, duplicate = ingest_file_bytes(
        db=db_session,
        content=content,
        original_file_name="race-copy.txt",
        source_type="OTHER",
    )

    assert duplicate is True
    assert duplicate_document.KnowledgeDocumentId == original_document.KnowledgeDocumentId
    assert db_session.query(KnowledgeDocument).count() == 1
    assert db_session.query(KnowledgeChunk).count() == original_document.ChunkCount


def test_docx_extraction_service_reads_generated_docx():
    doc = Document()
    doc.add_paragraph("Minerva extracts DOCX text.")
    doc.add_paragraph("The service remains read-only and advisory.")
    buffer = BytesIO()
    doc.save(buffer)

    text = extract_text_from_docx(buffer.getvalue())

    assert "Minerva extracts DOCX text." in text
    assert "read-only and advisory" in text


def test_json_upload_is_rejected_as_out_of_scope(client):
    response = client.post(
        "/api/v1/ingest/file",
        files={"file": ("evidence.json", b'{"payroll": true}', "application/json")},
        data={"source_type": "OTHER"},
    )

    assert response.status_code == 400
    assert "Operational JSON evidence ingestion is intentionally out of scope" in response.json()["detail"]


def test_invalid_source_type_is_rejected(client):
    response = client.post(
        "/api/v1/ingest/file",
        files={"file": ("invalid-source.txt", b"Invalid source type should be rejected.", "text/plain")},
        data={"source_type": "PAYROLL_RULES"},
    )

    assert response.status_code == 400
    assert "Invalid source_type" in response.json()["detail"]


def test_invalid_capability_status_is_rejected(client):
    response = client.post(
        "/api/v1/ingest/file",
        files={"file": ("invalid-status.txt", b"Invalid status should be rejected.", "text/plain")},
        data={"source_type": "OTHER", "capability_status": "DONE"},
    )

    assert response.status_code == 400
    assert "Invalid capability_status" in response.json()["detail"]
