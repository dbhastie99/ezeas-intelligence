from io import BytesIO
from pathlib import Path

from docx import Document


class UnsupportedDocumentTypeError(ValueError):
    pass


class EmptyDocumentError(ValueError):
    pass


JSON_OUT_OF_SCOPE_MESSAGE = (
    "Operational JSON evidence ingestion is intentionally out of scope for Minerva Walking Skeleton v0.1."
)


def extract_text_from_txt(content: bytes) -> str:
    return content.decode("utf-8-sig").strip()


def extract_text_from_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_text(content: bytes, file_name: str) -> str:
    extension = Path(file_name).suffix.lower()
    if extension == ".json":
        raise UnsupportedDocumentTypeError(JSON_OUT_OF_SCOPE_MESSAGE)
    if extension == ".txt":
        text = extract_text_from_txt(content)
    elif extension == ".docx":
        text = extract_text_from_docx(content)
    else:
        raise UnsupportedDocumentTypeError("Only TXT and DOCX files are supported in Minerva Walking Skeleton v0.1.")
    if not text:
        raise EmptyDocumentError("The uploaded document did not contain extractable text.")
    return text
